"""
run_benchmark.py — HyDE-bin, HyDE-qub, HyDE-con & HyGO Benchmark

Single unified run across 20 analytical benchmarks.
All four algorithms share the same evaluation budget and seed scheme.

Answers five research questions:
  (a) Mean final objective error across the 20 benchmark scenarios
      → Friedman test + Nemenyi post-hoc (block design) + per-scenario
        Kruskal-Wallis + Dunn's post-hoc (Bonferroni) + Cliff's delta
  (b) Convergence rate across 50 independent runs
      → Cochran's Q test + per-scenario chi-square / Fisher's exact test
  (c) Wall-clock cost per run
      → Friedman test on median wall times + per-scenario Kruskal-Wallis
        + speedup ratios
  (d) Practically meaningful margin by which each HyDE variant improves
      upon HyGO
      → Wilcoxon rank-sum + Cliff's delta + bootstrap 95% CI on mean
        difference
  (e) Performance as dimensionality scales from 2D to 25D across five
      shared scalable benchmark functions
      → Wilcoxon rank-sum + Cliff's delta + CV + degradation ratio

Cliff's delta thresholds (Robledo et al., 2025 / thesis methodology):
  negligible < 0.147, small >= 0.147, medium >= 0.33, large >= 0.474

Conditions: 5,000-evaluation budget, 50 independent runs, alpha = 0.05
"""

import csv
import json
import os
import time
import warnings
from itertools import combinations

import numpy as np
from scipy import stats
from scipy.stats import chi2_contingency, friedmanchisquare, kruskal, rankdata

from hyde_bench.benchmarks import FUNCTIONS, GLOBAL_OPTIMA


def _get_global_opt(fname, dim):
    """Resolve the global optimum for a given function and dimensionality."""
    val = GLOBAL_OPTIMA.get(fname)
    if val is not None:
        return float(val)
    if fname == 'styblinski_tang':
        return -39.16617 * dim
    return 0.0


# -- Configuration ----------------------------------------------------------

TEST_CASES = [
    ('ackley',          2),
    ('ackley',         25),
    ('beale',           2),
    ('booth',           2),
    ('bukin',           2),
    ('easom',           2),
    ('eggholder',       2),
    ('goldstein_price', 2),
    ('himmelblau',      2),
    ('holder_table',    2),
    ('levi',            2),
    ('matyas',          2),
    ('sphere',          2),
    ('sphere',         25),
    ('rastrigin',       2),
    ('rastrigin',      25),
    ('rosenbrock',      2),
    ('rosenbrock',     25),
    ('styblinski_tang', 2),
    ('styblinski_tang',25),
]

SCALABLE_FUNCTIONS = ['ackley', 'sphere', 'rastrigin', 'rosenbrock', 'styblinski_tang']

N_RUNS    = 50
MAX_EVALS = 50000
ALPHA     = 0.05

HERE      = os.path.dirname(os.path.abspath(__file__))
CHART_DIR = os.path.join(HERE, 'benchmark_charts')
CSV_DIR   = os.path.join(HERE, 'csv_data')
os.makedirs(CHART_DIR, exist_ok=True)
os.makedirs(CSV_DIR, exist_ok=True)

CLR_BIN  = '#f97316'
CLR_QUB  = '#38bdf8'
CLR_CON  = '#4ade80'
CLR_HYGO = '#a78bfa'
CLR_BG   = '#ffffff'
CLR_SURF = '#ffffff'
CLR_GRID = '#dddddd'
CLR_TEXT = '#000000'

ALGO_KEYS    = ['hyde_bin', 'hyde_qub', 'hyde_con', 'hygo']
ALGO_LABELS  = {
    'hyde_bin': 'HyDE-bin',
    'hyde_qub': 'HyDE-qub',
    'hyde_con': 'HyDE-con',
    'hygo':     'HyGO',
}
ALGO_COLOURS = {
    'hyde_bin': CLR_BIN,
    'hyde_qub': CLR_QUB,
    'hyde_con': CLR_CON,
    'hygo':     CLR_HYGO,
}
HYDE_KEYS = ['hyde_bin', 'hyde_qub', 'hyde_con']


# ============================================================================
# 1. RUNNER
# ============================================================================

def run_case(algo_class, fname, dim, algo_kwargs, n_runs, seed_base):
    func = FUNCTIONS[fname]
    results = []
    for i in range(n_runs):
        seed = seed_base + i * 1000 + dim * 7
        a = algo_class(func=func, fname=fname, dim=dim,
                       max_evals=MAX_EVALS, seed=seed, **algo_kwargs)
        t0 = time.perf_counter()
        r = a.run()
        r['wall_ms'] = (time.perf_counter() - t0) * 1000
        results.append(r)
    return results


def auc_normalized(curve):
    arr = np.array(curve, dtype=float)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", RuntimeWarning)
        arr = np.where(arr > 0, np.log10(arr + 1e-300), np.log10(1e-300))
    n = len(arr)
    _trap = getattr(np, 'trapezoid', None) or np.trapz
    return float(_trap(arr) / n) if n > 1 else float(arr[0])


def summarize(results, fname, dim):
    costs   = np.array([r['best_cost'] for r in results])
    evals   = np.array([r['evals']     for r in results])
    wall_ms = np.array([r['wall_ms']   for r in results])
    convs   = [r['conv_gen'] for r in results]
    n_conv  = sum(c is not None for c in convs)
    conv_gens = [c for c in convs if c is not None]

    histories = [r['gen_best'] for r in results]
    max_len   = max(len(h) for h in histories)
    padded    = [h + [h[-1]] * (max_len - len(h)) for h in histories]
    mean_curve = np.mean(padded, axis=0).tolist()
    aucs = [auc_normalized(h) for h in histories]

    global_opt = _get_global_opt(fname, dim)
    obj_errors = costs - global_opt

    mean_cost = float(np.mean(costs))
    std_cost  = float(np.std(costs, ddof=1))
    cv = (std_cost / abs(mean_cost)) if abs(mean_cost) > 1e-300 else 0.0

    # Per-run convergence binary vector (for Cochran's Q)
    conv_binary = [1 if c is not None else 0 for c in convs]

    return {
        'fname': fname, 'dim': dim, 'n_runs': len(results),
        'conv_pct':       round(100 * n_conv / len(results), 1),
        'conv_binary':    conv_binary,
        'mean_best':      mean_cost,
        'median_best':    float(np.median(costs)),
        'std_best':       std_cost,
        'min_best':       float(np.min(costs)),
        'max_best':       float(np.max(costs)),
        'iqr_best':       float(np.percentile(costs, 75) - np.percentile(costs, 25)),
        'cv':             cv,
        'mean_obj_error': float(np.mean(obj_errors)),
        'std_obj_error':  float(np.std(obj_errors, ddof=1)),
        'raw_obj_errors': obj_errors.tolist(),
        'mean_conv_gen':  round(float(np.mean(conv_gens)), 2) if conv_gens else None,
        'std_conv_gen':   round(float(np.std(conv_gens, ddof=1)), 2) if len(conv_gens) > 1 else None,
        'mean_auc':       float(np.mean(aucs)),
        'std_auc':        float(np.std(aucs, ddof=1)),
        'mean_evals':     float(np.mean(evals)),
        'std_evals':      float(np.std(evals, ddof=1)),
        'mean_wall_ms':   float(np.mean(wall_ms)),
        'std_wall_ms':    float(np.std(wall_ms, ddof=1)),
        'median_wall_ms': float(np.median(wall_ms)),
        'evals_per_ms':   float(np.mean(evals) / (np.mean(wall_ms) + 1e-9)),
        'raw_costs':      costs.tolist(),
        'raw_wall_ms':    wall_ms.tolist(),
        'raw_evals':      evals.tolist(),
        'raw_aucs':       aucs,
        'mean_curve':     mean_curve,
        'curves':         [r['gen_best'] for r in results[:8]],
    }


# ============================================================================
# 2. EFFECT SIZE — Cliff's delta
# ============================================================================

def cliffs_delta(a, b):
    """
    Cliff's delta: non-parametric effect size measure.
    Range: [-1, 1].  Negative d means group a tends to have smaller values.
    Thresholds (Robledo et al., 2025 / thesis):
      |d| < 0.147  negligible
      |d| >= 0.147 small
      |d| >= 0.33  medium
      |d| >= 0.474 large
    """
    a = np.asarray(a, dtype=float)
    b = np.asarray(b, dtype=float)
    n = len(a) * len(b)
    if n == 0:
        return 0.0
    more = np.sum(a[:, None] > b[None, :])
    less = np.sum(a[:, None] < b[None, :])
    return float((more - less) / n)


def interpret_cliffs_delta(d):
    ad = abs(d)
    if ad >= 0.474:
        return "large"
    elif ad >= 0.33:
        return "medium"
    elif ad >= 0.147:
        return "small"
    else:
        return "negligible"


def _is_effectively_constant(arr, rel_tol=1e-12, abs_tol=1e-30):
    rng = np.ptp(arr)
    if rng == 0:
        return True
    scale = max(abs(np.min(arr)), abs(np.max(arr)), 1e-300)
    return rng <= max(rel_tol * scale, abs_tol)


# ============================================================================
# 3. QUESTION (a) — Mean final objective error
#    Friedman test (block design) + Nemenyi post-hoc
#    Per-scenario: Kruskal-Wallis + Dunn's post-hoc (Bonferroni) + Cliff's delta
# ============================================================================

def dunns_posthoc_bonferroni(groups):
    """
    Dunn's post-hoc test with Bonferroni correction for pairwise comparisons
    following a significant Kruskal-Wallis test.
    """
    labels = list(groups.keys())
    arrays = [np.asarray(groups[l], dtype=float) for l in labels]
    k = len(labels)
    n_pairs = k * (k - 1) // 2

    all_vals = np.concatenate(arrays)
    N = len(all_vals)
    all_ranks = rankdata(all_vals)

    sizes = [len(a) for a in arrays]
    mean_ranks = []
    idx = 0
    for sz in sizes:
        mean_ranks.append(np.mean(all_ranks[idx:idx + sz]))
        idx += sz

    _, tie_counts = np.unique(all_vals, return_counts=True)
    tie_correction = np.sum([tc ** 3 - tc for tc in tie_counts if tc > 1])
    base_var = (N * (N + 1) / 12.0) - (tie_correction / (12.0 * (N - 1))) if N > 1 else 1.0

    rows = []
    for (i, j) in combinations(range(k), 2):
        diff = mean_ranks[i] - mean_ranks[j]
        se2 = base_var * (1.0 / sizes[i] + 1.0 / sizes[j])
        se = np.sqrt(max(se2, 1e-300))
        z = diff / se
        p_raw = float(2.0 * (1.0 - stats.norm.cdf(abs(z))))
        p_adj = min(p_raw * n_pairs, 1.0)

        d = cliffs_delta(arrays[i], arrays[j])

        direction = (f"{labels[i]} better" if diff < 0 else
                     f"{labels[j]} better" if diff > 0 else "equal")
        rows.append({
            'pair': f"{labels[i]} vs {labels[j]}",
            'label_i': labels[i], 'label_j': labels[j],
            'mean_rank_i': float(mean_ranks[i]),
            'mean_rank_j': float(mean_ranks[j]),
            'mean_rank_diff': float(diff),
            'z_stat': float(z),
            'p_raw': float(p_raw),
            'p_bonferroni': float(p_adj),
            'significant': bool(p_adj < ALPHA),
            'cliffs_delta': float(d),
            'd_magnitude': interpret_cliffs_delta(d),
            'direction': direction,
        })
    return rows


def run_kruskal_per_scenario(key, entry):
    """
    Per-scenario Kruskal-Wallis + Dunn's post-hoc across all 4 algorithms.
    """
    groups = {k: np.array(entry[k]['raw_costs']) for k in ALGO_KEYS}
    arrays = list(groups.values())
    all_values = np.concatenate(arrays)

    if _is_effectively_constant(all_values):
        means = {k: entry[k]['mean_best'] for k in ALGO_KEYS}
        best_algo = min(means, key=means.get)
        return {
            'key': key,
            'h_stat': 0.0, 'p_kruskal': 1.0, 'sig': False,
            'best_algo': best_algo,
            'posthoc': dunns_posthoc_bonferroni(groups),
            'note': 'All groups effectively constant.',
        }

    h_stat, p_kruskal = kruskal(*arrays)
    if np.isnan(h_stat) or np.isnan(p_kruskal):
        h_stat, p_kruskal = 0.0, 1.0

    posthoc = dunns_posthoc_bonferroni(groups)
    means = {k: entry[k]['mean_best'] for k in ALGO_KEYS}
    best_algo = min(means, key=means.get)

    return {
        'key': key,
        'h_stat': float(h_stat),
        'p_kruskal': float(p_kruskal),
        'sig': bool(p_kruskal < ALPHA),
        'best_algo': best_algo,
        'posthoc': posthoc,
    }


def friedman_objective_error(all_results):
    """
    Friedman test on mean objective error across benchmarks (block design).
    Each benchmark is a block; the 4 algorithms are repeated measures.
    The observation per cell is the mean objective error over 50 runs.

    Followed by Nemenyi post-hoc test for pairwise comparisons.
    """
    bench_keys = list(all_results.keys())
    n_benchmarks = len(bench_keys)

    # Build matrix: rows = benchmarks, columns = algorithms
    data = np.zeros((n_benchmarks, len(ALGO_KEYS)))
    for i, bk in enumerate(bench_keys):
        for j, ak in enumerate(ALGO_KEYS):
            data[i, j] = all_results[bk][ak]['mean_obj_error']

    # Friedman test
    chi2, p_friedman = friedmanchisquare(*[data[:, j] for j in range(len(ALGO_KEYS))])
    if np.isnan(chi2):
        chi2, p_friedman = 0.0, 1.0

    # Compute mean ranks across benchmarks for each algorithm
    ranks = np.zeros_like(data)
    for i in range(n_benchmarks):
        ranks[i] = rankdata(data[i])
    mean_ranks = ranks.mean(axis=0)

    # Nemenyi post-hoc: critical difference
    k = len(ALGO_KEYS)
    # q_alpha values for Nemenyi test (from studentized range table / k groups)
    # For k=4, alpha=0.05: q_0.05 ≈ 2.569
    q_alpha_table = {3: 2.343, 4: 2.569, 5: 2.728, 6: 2.850}
    q_alpha = q_alpha_table.get(k, 2.569)
    cd = q_alpha * np.sqrt(k * (k + 1) / (6.0 * n_benchmarks))

    nemenyi_pairs = []
    for (i, j) in combinations(range(k), 2):
        diff = abs(mean_ranks[i] - mean_ranks[j])
        sig = diff > cd
        if mean_ranks[i] < mean_ranks[j]:
            direction = f"{ALGO_LABELS[ALGO_KEYS[i]]} better"
        elif mean_ranks[j] < mean_ranks[i]:
            direction = f"{ALGO_LABELS[ALGO_KEYS[j]]} better"
        else:
            direction = "equal"
        nemenyi_pairs.append({
            'pair': f"{ALGO_LABELS[ALGO_KEYS[i]]} vs {ALGO_LABELS[ALGO_KEYS[j]]}",
            'mean_rank_i': float(mean_ranks[i]),
            'mean_rank_j': float(mean_ranks[j]),
            'rank_diff': float(diff),
            'critical_diff': float(cd),
            'significant': bool(sig),
            'direction': direction,
        })

    best_idx = int(np.argmin(mean_ranks))
    return {
        'chi2': float(chi2),
        'p_friedman': float(p_friedman),
        'sig': bool(p_friedman < ALPHA),
        'mean_ranks': {ALGO_KEYS[j]: float(mean_ranks[j]) for j in range(k)},
        'best_algo': ALGO_KEYS[best_idx],
        'critical_diff': float(cd),
        'nemenyi': nemenyi_pairs,
        'n_benchmarks': n_benchmarks,
    }


# ============================================================================
# 4. QUESTION (b) — Convergence rate
#    Cochran's Q test + per-scenario chi-square
# ============================================================================

def cochrans_q_test(all_results):
    """
    Cochran's Q test: are convergence rates significantly different across
    the 4 algorithms?  Each benchmark contributes N_RUNS binary observations
    per algorithm (converged=1, not=0).

    Q = (k-1) * [k * sum(T_j^2) - T..^2] / [k * T.. - sum(L_i^2)]
    where T_j = column total for algorithm j, L_i = row total for run i,
    T.. = grand total.
    """
    bench_keys = list(all_results.keys())
    k = len(ALGO_KEYS)

    # Stack all binary convergence vectors: rows = (benchmark, run), cols = algorithms
    all_rows = []
    for bk in bench_keys:
        entry = all_results[bk]
        n = len(entry[ALGO_KEYS[0]]['conv_binary'])
        for run_i in range(n):
            row = [entry[ak]['conv_binary'][run_i] for ak in ALGO_KEYS]
            all_rows.append(row)

    mat = np.array(all_rows, dtype=float)  # shape (N_total, k)
    N = mat.shape[0]

    T_j = mat.sum(axis=0)       # column totals
    L_i = mat.sum(axis=1)       # row totals
    T_dot = mat.sum()

    numerator = (k - 1) * (k * np.sum(T_j ** 2) - T_dot ** 2)
    denominator = k * T_dot - np.sum(L_i ** 2)

    if denominator == 0:
        Q_stat, p_val = 0.0, 1.0
    else:
        Q_stat = float(numerator / denominator)
        p_val = float(1.0 - stats.chi2.cdf(Q_stat, df=k - 1))

    conv_rates = {ak: float(T_j[i] / N * 100) for i, ak in enumerate(ALGO_KEYS)}

    return {
        'Q_stat': Q_stat,
        'p_value': p_val,
        'df': k - 1,
        'sig': bool(p_val < ALPHA),
        'n_total_obs': N,
        'conv_rates': conv_rates,
    }


def chi2_convergence_per_scenario(key, entry):
    """
    Per-scenario chi-square test on the 4x2 contingency table
    (converged vs not-converged for each algorithm).
    Falls back to Fisher's exact test for 2x2 sub-comparisons.
    """
    n_runs = entry[ALGO_KEYS[0]]['n_runs']
    conv_counts = []
    for ak in ALGO_KEYS:
        n_conv = sum(entry[ak]['conv_binary'])
        conv_counts.append([n_conv, n_runs - n_conv])

    table = np.array(conv_counts)  # shape (4, 2)

    # If all identical convergence rates, skip
    if np.all(table[:, 0] == table[0, 0]):
        return {
            'key': key, 'chi2': 0.0, 'p_value': 1.0, 'sig': False,
            'conv_counts': {ak: int(table[i, 0]) for i, ak in enumerate(ALGO_KEYS)},
            'note': 'All algorithms have identical convergence counts.',
        }

    # Check for zero rows/columns that would make chi2 degenerate
    if np.any(table.sum(axis=0) == 0):
        return {
            'key': key, 'chi2': 0.0, 'p_value': 1.0, 'sig': False,
            'conv_counts': {ak: int(table[i, 0]) for i, ak in enumerate(ALGO_KEYS)},
            'note': 'Degenerate table (all or none converged).',
        }

    chi2_stat, p_val, dof, _ = chi2_contingency(table)
    if np.isnan(chi2_stat):
        chi2_stat, p_val = 0.0, 1.0

    return {
        'key': key,
        'chi2': float(chi2_stat),
        'p_value': float(p_val),
        'dof': int(dof),
        'sig': bool(p_val < ALPHA),
        'conv_counts': {ak: int(table[i, 0]) for i, ak in enumerate(ALGO_KEYS)},
    }


# ============================================================================
# 5. QUESTION (c) — Wall-clock cost per run
#    Friedman test on median wall times + per-scenario Kruskal-Wallis
# ============================================================================

def friedman_wall_time(all_results):
    """
    Friedman test on median wall-clock time per algorithm across benchmarks.
    """
    bench_keys = list(all_results.keys())
    n = len(bench_keys)

    data = np.zeros((n, len(ALGO_KEYS)))
    for i, bk in enumerate(bench_keys):
        for j, ak in enumerate(ALGO_KEYS):
            data[i, j] = all_results[bk][ak]['median_wall_ms']

    chi2, p_val = friedmanchisquare(*[data[:, j] for j in range(len(ALGO_KEYS))])
    if np.isnan(chi2):
        chi2, p_val = 0.0, 1.0

    ranks = np.zeros_like(data)
    for i in range(n):
        ranks[i] = rankdata(data[i])
    mean_ranks = ranks.mean(axis=0)

    # Speedup ratios relative to HyGO
    hygo_idx = ALGO_KEYS.index('hygo')
    grand_means = {ak: float(np.mean([all_results[bk][ak]['mean_wall_ms']
                                       for bk in bench_keys]))
                   for ak in ALGO_KEYS}
    hygo_mean = grand_means['hygo']
    speedup_vs_hygo = {ak: hygo_mean / max(grand_means[ak], 1e-9)
                       for ak in ALGO_KEYS}

    fastest = min(grand_means, key=grand_means.get)

    return {
        'chi2': float(chi2),
        'p_friedman': float(p_val),
        'sig': bool(p_val < ALPHA),
        'mean_ranks': {ALGO_KEYS[j]: float(mean_ranks[j]) for j in range(len(ALGO_KEYS))},
        'grand_means_ms': grand_means,
        'speedup_vs_hygo': speedup_vs_hygo,
        'fastest': fastest,
    }


def kruskal_wall_time_per_scenario(key, entry):
    """Per-scenario Kruskal-Wallis on raw wall-clock times."""
    groups = [np.array(entry[ak]['raw_wall_ms']) for ak in ALGO_KEYS]
    all_vals = np.concatenate(groups)

    if _is_effectively_constant(all_vals):
        return {'key': key, 'h_stat': 0.0, 'p_value': 1.0, 'sig': False}

    h_stat, p_val = kruskal(*groups)
    if np.isnan(h_stat):
        h_stat, p_val = 0.0, 1.0

    means = {ak: entry[ak]['mean_wall_ms'] for ak in ALGO_KEYS}
    fastest = min(means, key=means.get)

    return {
        'key': key,
        'h_stat': float(h_stat),
        'p_value': float(p_val),
        'sig': bool(p_val < ALPHA),
        'means_ms': means,
        'fastest': fastest,
    }


# ============================================================================
# 6. QUESTION (d) — Practically meaningful margin vs HyGO
#    Wilcoxon rank-sum + Cliff's delta + bootstrap 95% CI
# ============================================================================

def bootstrap_mean_diff_ci(a, b, n_boot=10000, ci=0.95, seed=42):
    """
    Bootstrap 95% CI on the mean difference (a - b).
    Negative CI means 'a' tends to be smaller (better for minimization).
    """
    rng = np.random.default_rng(seed)
    a = np.asarray(a, dtype=float)
    b = np.asarray(b, dtype=float)
    diffs = np.empty(n_boot)
    na, nb = len(a), len(b)
    for i in range(n_boot):
        sa = a[rng.integers(0, na, na)]
        sb = b[rng.integers(0, nb, nb)]
        diffs[i] = sa.mean() - sb.mean()
    alpha_half = (1 - ci) / 2
    lo = float(np.percentile(diffs, 100 * alpha_half))
    hi = float(np.percentile(diffs, 100 * (1 - alpha_half)))
    return {
        'mean_diff': float(np.mean(diffs)),
        'ci_lo': lo,
        'ci_hi': hi,
        'ci_level': ci,
    }


def wilcoxon_margin_vs_hygo(key, entry, hyde_key):
    """
    Wilcoxon rank-sum + Cliff's delta + bootstrap CI for one HyDE vs HyGO.
    """
    hyde_costs = np.array(entry[hyde_key]['raw_costs'])
    hygo_costs = np.array(entry['hygo']['raw_costs'])

    hyde_mean = float(np.mean(hyde_costs))
    hygo_mean = float(np.mean(hygo_costs))

    hyde_const = _is_effectively_constant(hyde_costs)
    hygo_const = _is_effectively_constant(hygo_costs)

    note = None
    if hyde_const and hygo_const:
        u_stat, p_val = 0.0, 1.0
        note = 'Both groups effectively constant.'
    elif hyde_const or hygo_const:
        u_stat = 0.0
        p_val = 1.0 if hyde_mean == hygo_mean else 0.0
        note = 'One group effectively constant.'
    else:
        u_stat, p_val = stats.mannwhitneyu(hyde_costs, hygo_costs,
                                            alternative='two-sided')
        if np.isnan(u_stat):
            u_stat = 0.0
        if np.isnan(p_val):
            p_val = 1.0

    d = cliffs_delta(hyde_costs, hygo_costs)
    boot = bootstrap_mean_diff_ci(hyde_costs, hygo_costs)

    if hyde_mean < hygo_mean:
        direction = f"{ALGO_LABELS[hyde_key]} better"
    elif hygo_mean < hyde_mean:
        direction = "HyGO better"
    else:
        direction = "equal"

    result = {
        'key': key,
        'hyde_key': hyde_key,
        'hyde_label': ALGO_LABELS[hyde_key],
        'u_stat': float(u_stat),
        'p_value': float(p_val),
        'sig': bool(p_val < ALPHA),
        'hyde_mean': hyde_mean,
        'hygo_mean': hygo_mean,
        'mean_diff': hyde_mean - hygo_mean,
        'cliffs_delta': float(d),
        'd_magnitude': interpret_cliffs_delta(d),
        'bootstrap_ci_lo': boot['ci_lo'],
        'bootstrap_ci_hi': boot['ci_hi'],
        'direction': direction,
    }
    if note:
        result['note'] = note
    return result


# ============================================================================
# 7. QUESTION (e) — Dimensionality scaling 2D → 25D
#    Wilcoxon + Cliff's delta + CV + degradation ratio
# ============================================================================

def run_scaling_analysis(all_results):
    """
    For each scalable function and each algorithm:
    - Wilcoxon rank-sum (2D vs 25D raw costs)
    - Cliff's delta
    - CV at 2D and 25D
    - Degradation ratio: mean_25D / mean_2D
    """
    rows = []
    for fname in SCALABLE_FUNCTIONS:
        key_2d  = f"{fname}_2D"
        key_25d = f"{fname}_25D"
        if key_2d not in all_results or key_25d not in all_results:
            continue

        for ak in ALGO_KEYS:
            costs_2d  = np.array(all_results[key_2d][ak]['raw_costs'])
            costs_25d = np.array(all_results[key_25d][ak]['raw_costs'])

            cv_2d  = all_results[key_2d][ak]['cv']
            cv_25d = all_results[key_25d][ak]['cv']

            const_2d  = _is_effectively_constant(costs_2d)
            const_25d = _is_effectively_constant(costs_25d)

            note = None
            if const_2d and const_25d:
                u_stat, p_val = 0.0, 1.0
                note = 'Both groups effectively constant.'
            elif const_2d or const_25d:
                u_stat, p_val = 0.0, 1.0
                note = 'One group effectively constant.'
            else:
                u_stat, p_val = stats.mannwhitneyu(costs_2d, costs_25d,
                                                    alternative='two-sided')
                if np.isnan(u_stat):
                    u_stat = 0.0
                if np.isnan(p_val):
                    p_val = 1.0

            d = cliffs_delta(costs_2d, costs_25d)
            mean_2d  = float(np.mean(costs_2d))
            mean_25d = float(np.mean(costs_25d))

            # Degradation ratio
            if abs(mean_2d) > 1e-300:
                deg_ratio = mean_25d / mean_2d
            else:
                deg_ratio = float('inf') if abs(mean_25d) > 1e-300 else 1.0

            row = {
                'fname': fname,
                'algo_key': ak,
                'algo_label': ALGO_LABELS[ak],
                'u_stat': float(u_stat),
                'p_value': float(p_val),
                'sig': bool(p_val < ALPHA),
                'mean_2d': mean_2d,
                'mean_25d': mean_25d,
                'cv_2d': float(cv_2d),
                'cv_25d': float(cv_25d),
                'cliffs_delta': float(d),
                'd_magnitude': interpret_cliffs_delta(d),
                'degradation_ratio': float(deg_ratio),
                'direction': ('degraded' if mean_25d > mean_2d else
                              'improved' if mean_25d < mean_2d else 'equal'),
            }
            if note:
                row['note'] = note
            rows.append(row)

    return rows


# ============================================================================
# 8. FIGURE 5 CONVERGENCE CURVES
# ============================================================================

FIGURE5_PAIRS = [
    ('rosenbrock', 2), ('rosenbrock', 25),
    ('rastrigin',  2), ('rastrigin',  25),
    ('sphere',     2), ('sphere',     25),
]


def make_figure5_curves(all_results):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        print("  [skip Figure 5 curves — matplotlib not installed]")
        return

    plt.rcParams.update({
        'figure.facecolor': CLR_BG, 'axes.facecolor': CLR_SURF,
        'axes.edgecolor': CLR_GRID, 'axes.labelcolor': CLR_TEXT,
        'text.color': CLR_TEXT, 'xtick.color': CLR_TEXT,
        'ytick.color': CLR_TEXT, 'grid.color': CLR_GRID,
        'font.size': 9,
    })

    valid_pairs = [(fn, d) for fn, d in FIGURE5_PAIRS
                   if f"{fn}_{d}D" in all_results]
    if not valid_pairs:
        return

    fig, axes = plt.subplots(2, 3, figsize=(15, 9))
    axes = axes.flatten()

    for ax, (fname, dim) in zip(axes, valid_pairs):
        key   = f"{fname}_{dim}D"
        entry = all_results[key]
        for ak in ALGO_KEYS:
            curve = entry[ak]['mean_curve']
            ax.plot(curve, label=ALGO_LABELS[ak],
                    color=ALGO_COLOURS[ak], linewidth=2)
        ax.set_title(f"{fname.capitalize()} {dim}D", fontsize=10, fontweight='bold')
        ax.set_xlabel('Generation')
        ax.set_ylabel(f'Mean Best Cost ({N_RUNS} runs)')
        if any(v > 0 for v in entry[ALGO_KEYS[0]]['mean_curve']):
            ax.set_yscale('symlog', linthresh=1e-10)
        ax.legend(fontsize=7)
        ax.grid(True, alpha=0.3)

    for ax in axes[len(valid_pairs):]:
        ax.set_visible(False)

    fig.suptitle(
        'Figure 5 — Mean Convergence Curves: Rosenbrock, Rastrigin, Sphere (2D & 25D)',
        fontsize=12, fontweight='bold')
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'figure5_convergence_curves.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)
    print("  Figure 5 curves saved.")


# ============================================================================
# 9. CHARTING
# ============================================================================

def make_charts(all_results, kruskal_results, margin_results):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        print("  [skip charts — matplotlib not installed]")
        return

    plt.rcParams.update({
        'figure.facecolor': CLR_BG, 'axes.facecolor': CLR_SURF,
        'axes.edgecolor': CLR_GRID, 'axes.labelcolor': CLR_TEXT,
        'text.color': CLR_TEXT, 'xtick.color': CLR_TEXT,
        'ytick.color': CLR_TEXT, 'grid.color': CLR_GRID,
        'font.size': 9,
    })

    # --- Per-benchmark: convergence + box plot ---
    for bench_key, entry in all_results.items():
        fname = entry['hyde_bin']['fname']
        dim   = entry['hyde_bin']['dim']

        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))
        fig.suptitle(f'{fname.upper()} {dim}D — HyDE variants vs HyGO',
                     fontsize=13, fontweight='bold')

        all_negative = True
        for ak in ALGO_KEYS:
            curve = entry[ak]['mean_curve']
            ax1.plot(curve, label=ALGO_LABELS[ak],
                     color=ALGO_COLOURS[ak], linewidth=2)
            if any(v > 0 for v in curve):
                all_negative = False
        ax1.set_xlabel('Generation')
        ax1.set_ylabel(f'Best Cost (mean of {N_RUNS} runs)')
        if not all_negative:
            ax1.set_yscale('symlog', linthresh=1e-10)
        ax1.legend()
        ax1.grid(True, alpha=0.3)
        ax1.set_title('Mean Convergence Curve')

        data   = [entry[ak]['raw_costs'] for ak in ALGO_KEYS]
        labels = [ALGO_LABELS[ak]        for ak in ALGO_KEYS]
        bp = ax2.boxplot(data, tick_labels=labels, patch_artist=True)
        for patch, ak in zip(bp['boxes'], ALGO_KEYS):
            patch.set_facecolor(ALGO_COLOURS[ak])
            patch.set_alpha(0.7)
        for element in ['whiskers', 'caps', 'medians']:
            for line in bp[element]:
                line.set_color(CLR_TEXT)
        ax2.set_ylabel('Best Cost')
        ax2.set_title(f'Final Cost Distribution ({N_RUNS} runs)')
        ax2.ticklabel_format(axis='y', useOffset=False)
        ax2.grid(True, alpha=0.3)

        fig.tight_layout()
        fig.savefig(os.path.join(CHART_DIR, f'{bench_key}.png'),
                    dpi=150, bbox_inches='tight')
        plt.close(fig)

    # --- (d) Margin vs HyGO: wins bar chart ---
    wins = {k: 0 for k in HYDE_KEYS}
    wins['hygo'] = 0
    ties = 0
    for pr in margin_results:
        if pr['sig']:
            if 'HyGO better' in pr['direction']:
                wins['hygo'] += 1
            else:
                wins[pr['hyde_key']] += 1
        else:
            ties += 1

    fig, ax = plt.subplots(figsize=(10, 5))
    bar_labels = [ALGO_LABELS[k] for k in HYDE_KEYS] + ['HyGO', 'No Sig. Diff.']
    vals   = [wins[k] for k in HYDE_KEYS] + [wins['hygo'], ties]
    colors = [ALGO_COLOURS[k] for k in HYDE_KEYS] + [CLR_HYGO, CLR_GRID]
    bars = ax.bar(bar_labels, vals, color=colors, alpha=0.8)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                str(v), ha='center', fontweight='bold', color=CLR_TEXT)
    ax.set_ylabel('Number of Benchmarks')
    ax.set_title(f"(d) HyDE vs HyGO — Wilcoxon Significant Wins (alpha = {ALPHA})",
                 fontweight='bold')
    ax.grid(True, alpha=0.3, axis='y')
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'qd_margin_wins.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)

    # --- (a) Best variant wins bar chart ---
    kruskal_wins = {k: 0 for k in ALGO_KEYS}
    for kr in kruskal_results:
        best = kr['best_algo']
        if best in kruskal_wins:
            kruskal_wins[best] += 1

    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.bar([ALGO_LABELS[k] for k in ALGO_KEYS],
                  [kruskal_wins[k] for k in ALGO_KEYS],
                  color=[ALGO_COLOURS[k] for k in ALGO_KEYS], alpha=0.8)
    for bar, k in zip(bars, ALGO_KEYS):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                str(kruskal_wins[k]), ha='center', fontweight='bold', color=CLR_TEXT)
    ax.set_ylabel('Number of Wins (by mean best cost)')
    ax.set_title('(a) Best Algorithm per Scenario — by Mean Objective Error',
                 fontweight='bold')
    ax.grid(True, alpha=0.3, axis='y')
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'qa_objective_error_wins.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)


def make_scaling_chart(scaling_results):
    """(e) Bar chart of CV at 25D and degradation ratio heatmap."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        return

    plt.rcParams.update({
        'figure.facecolor': CLR_BG, 'axes.facecolor': CLR_SURF,
        'axes.edgecolor': CLR_GRID, 'axes.labelcolor': CLR_TEXT,
        'text.color': CLR_TEXT, 'xtick.color': CLR_TEXT,
        'ytick.color': CLR_TEXT, 'grid.color': CLR_GRID,
        'font.size': 8,
    })

    fnames_present = list(dict.fromkeys(r['fname'] for r in scaling_results))
    n_fn  = len(fnames_present)
    n_algo = len(ALGO_KEYS)
    x = np.arange(n_fn)
    width = 0.2

    # CV at 25D bar chart
    fig, ax = plt.subplots(figsize=(max(10, n_fn * 1.5), 5))
    for i, ak in enumerate(ALGO_KEYS):
        cv25 = []
        for fn in fnames_present:
            row = next((r for r in scaling_results
                        if r['fname'] == fn and r['algo_key'] == ak), None)
            cv25.append(row['cv_25d'] if row else 0.0)
        offset = (i - n_algo / 2 + 0.5) * width
        ax.bar(x + offset, cv25, width,
               label=ALGO_LABELS[ak], color=ALGO_COLOURS[ak], alpha=0.85)

    ax.set_xticks(x)
    ax.set_xticklabels([fn.replace('_', '\n') for fn in fnames_present], fontsize=8)
    ax.set_ylabel('Coefficient of Variation at 25D')
    ax.set_title('(e) Result Consistency at 25D — CV per Benchmark',
                 fontweight='bold')
    ax.legend(loc='upper right')
    ax.grid(True, alpha=0.3, axis='y')
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'qe_cv_25d.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)

    # Degradation ratio heatmap
    fig, ax = plt.subplots(figsize=(max(8, n_fn * 1.2), 4))
    deg_matrix = np.zeros((n_algo, n_fn))
    for j, fn in enumerate(fnames_present):
        for i, ak in enumerate(ALGO_KEYS):
            row = next((r for r in scaling_results
                        if r['fname'] == fn and r['algo_key'] == ak), None)
            deg_matrix[i, j] = row['degradation_ratio'] if row else 1.0

    # Clip for display
    deg_display = np.clip(deg_matrix, 0, np.percentile(deg_matrix, 95) * 1.2)

    im = ax.imshow(deg_display, aspect='auto', cmap='YlOrRd')
    ax.set_xticks(range(n_fn))
    ax.set_xticklabels([fn.replace('_', '\n') for fn in fnames_present], fontsize=8)
    ax.set_yticks(range(n_algo))
    ax.set_yticklabels([ALGO_LABELS[ak] for ak in ALGO_KEYS], fontsize=9)
    for i in range(n_algo):
        for j in range(n_fn):
            val = deg_matrix[i, j]
            txt = f"{val:.1f}" if abs(val) < 1e6 else f"{val:.1e}"
            ax.text(j, i, txt, ha='center', va='center', fontsize=7,
                    color='white' if deg_display[i, j] > deg_display.max() * 0.6 else 'black')
    fig.colorbar(im, ax=ax, label='Degradation Ratio (mean_25D / mean_2D)')
    ax.set_title('(e) Degradation Ratio: 2D → 25D (lower = scales better)',
                 fontweight='bold')
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'qe_degradation_heatmap.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)


def make_cost_charts(all_results):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        return

    plt.rcParams.update({
        'figure.facecolor': CLR_BG, 'axes.facecolor': CLR_SURF,
        'axes.edgecolor': CLR_GRID, 'axes.labelcolor': CLR_TEXT,
        'text.color': CLR_TEXT, 'xtick.color': CLR_TEXT,
        'ytick.color': CLR_TEXT, 'grid.color': CLR_GRID,
        'font.size': 8,
    })

    bench_keys = list(all_results.keys())
    n_bench    = len(bench_keys)
    n_algo     = len(ALGO_KEYS)
    x          = np.arange(n_bench)
    width      = 0.2

    fig, ax = plt.subplots(figsize=(max(14, n_bench * 0.8), 6))
    for i, ak in enumerate(ALGO_KEYS):
        times  = [all_results[bk][ak]['mean_wall_ms'] for bk in bench_keys]
        offset = (i - n_algo / 2 + 0.5) * width
        ax.bar(x + offset, times, width,
               label=ALGO_LABELS[ak], color=ALGO_COLOURS[ak], alpha=0.85)
    ax.set_xticks(x)
    ax.set_xticklabels([bk.replace('_', '\n') for bk in bench_keys],
                       rotation=0, fontsize=7)
    ax.set_ylabel('Mean Wall-Clock Time (ms)')
    ax.set_title('(c) Computational Cost — Mean Wall-Clock Time per Benchmark',
                 fontweight='bold')
    ax.legend(loc='upper left')
    ax.grid(True, alpha=0.3, axis='y')
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'qc_wall_time_per_benchmark.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(8, 5))
    means = [np.mean([all_results[bk][ak]['mean_wall_ms'] for bk in bench_keys])
             for ak in ALGO_KEYS]
    stds  = [np.std([all_results[bk][ak]['mean_wall_ms'] for bk in bench_keys])
             for ak in ALGO_KEYS]
    bars = ax.bar([ALGO_LABELS[k] for k in ALGO_KEYS], means,
                  yerr=stds, capsize=5,
                  color=[ALGO_COLOURS[k] for k in ALGO_KEYS], alpha=0.85)
    for bar, m in zip(bars, means):
        ax.text(bar.get_x() + bar.get_width() / 2,
                bar.get_height() + max(stds) * 0.05,
                f'{m:.0f}ms', ha='center', fontsize=9,
                fontweight='bold', color=CLR_TEXT)
    ax.set_ylabel('Mean Wall-Clock Time (ms)')
    ax.set_title('(c) Computational Cost — Average Across All Benchmarks',
                 fontweight='bold')
    ax.grid(True, alpha=0.3, axis='y')
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'qc_wall_time_summary.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)


def make_convergence_charts(all_results):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        return

    plt.rcParams.update({
        'figure.facecolor': CLR_BG, 'axes.facecolor': CLR_SURF,
        'axes.edgecolor': CLR_GRID, 'axes.labelcolor': CLR_TEXT,
        'text.color': CLR_TEXT, 'xtick.color': CLR_TEXT,
        'ytick.color': CLR_TEXT, 'grid.color': CLR_GRID,
        'font.size': 8,
    })

    bench_keys = list(all_results.keys())
    n_bench    = len(bench_keys)
    n_algo     = len(ALGO_KEYS)
    x          = np.arange(n_bench)
    width      = 0.2

    fig, ax = plt.subplots(figsize=(max(14, n_bench * 0.8), 6))
    for i, ak in enumerate(ALGO_KEYS):
        rates  = [all_results[bk][ak]['conv_pct'] for bk in bench_keys]
        offset = (i - n_algo / 2 + 0.5) * width
        ax.bar(x + offset, rates, width,
               label=ALGO_LABELS[ak], color=ALGO_COLOURS[ak], alpha=0.85)
    ax.set_xticks(x)
    ax.set_xticklabels([bk.replace('_', '\n') for bk in bench_keys],
                       rotation=0, fontsize=7)
    ax.set_ylim(0, 110)
    ax.set_ylabel('Convergence Rate (%)')
    ax.set_title(f'(b) Convergence Rate — % of {N_RUNS} Runs that Converged per Benchmark',
                 fontweight='bold')
    ax.legend(loc='upper left')
    ax.grid(True, alpha=0.3, axis='y')
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'qb_convergence_rate.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7, 6))
    for ak in ALGO_KEYS:
        pcts  = [all_results[bk][ak]['conv_pct']          for bk in bench_keys]
        gens  = [all_results[bk][ak].get('mean_conv_gen') for bk in bench_keys]
        valid = [(p, g) for p, g in zip(pcts, gens) if g is not None]
        if not valid:
            continue
        avg_pct = np.mean([v[0] for v in valid])
        avg_gen = np.mean([v[1] for v in valid])
        ax.scatter(avg_gen, avg_pct, s=180, color=ALGO_COLOURS[ak],
                   label=ALGO_LABELS[ak], zorder=3)
        ax.annotate(ALGO_LABELS[ak], xy=(avg_gen, avg_pct),
                    xytext=(6, 4), textcoords='offset points',
                    color=ALGO_COLOURS[ak], fontsize=9, fontweight='bold')
    ax.set_xlabel('Mean Generation at First Convergence  (lower = faster)')
    ax.set_ylabel('Mean Convergence Rate %  (higher = more reliable)')
    ax.set_title('(b) Convergence Summary\nTop-left = fastest and most reliable',
                 fontweight='bold')
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    fig.savefig(os.path.join(CHART_DIR, 'qb_convergence_summary.png'),
                dpi=150, bbox_inches='tight')
    plt.close(fig)


def make_bootstrap_ci_chart(margin_results):
    """(d) Forest plot of bootstrap CIs on mean difference (HyDE - HyGO)."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        return

    plt.rcParams.update({
        'figure.facecolor': CLR_BG, 'axes.facecolor': CLR_SURF,
        'axes.edgecolor': CLR_GRID, 'axes.labelcolor': CLR_TEXT,
        'text.color': CLR_TEXT, 'xtick.color': CLR_TEXT,
        'ytick.color': CLR_TEXT, 'grid.color': CLR_GRID,
        'font.size': 8,
    })

    for hk in HYDE_KEYS:
        hk_results = [r for r in margin_results if r['hyde_key'] == hk]
        if not hk_results:
            continue

        fig, ax = plt.subplots(figsize=(10, max(4, len(hk_results) * 0.35)))
        y_pos = range(len(hk_results))
        means = [r['mean_diff'] for r in hk_results]
        ci_lo = [r['bootstrap_ci_lo'] for r in hk_results]
        ci_hi = [r['bootstrap_ci_hi'] for r in hk_results]
        labels = [r['key'] for r in hk_results]

        xerr_lo = [m - lo for m, lo in zip(means, ci_lo)]
        xerr_hi = [hi - m for m, hi in zip(means, ci_hi)]

        colors = ['green' if hi < 0 else 'red' if lo > 0 else 'gray'
                  for lo, hi in zip(ci_lo, ci_hi)]

        ax.barh(y_pos, means, xerr=[xerr_lo, xerr_hi],
                color=colors, alpha=0.7, capsize=3, height=0.6)
        ax.axvline(0, color='black', linewidth=0.8, linestyle='--')
        ax.set_yticks(y_pos)
        ax.set_yticklabels(labels, fontsize=7)
        ax.set_xlabel('Mean Difference (HyDE - HyGO)\n← HyDE better | HyGO better →')
        ax.set_title(f'(d) {ALGO_LABELS[hk]} vs HyGO — Bootstrap 95% CI on Mean Difference',
                     fontweight='bold')
        ax.grid(True, alpha=0.3, axis='x')
        fig.tight_layout()
        fig.savefig(os.path.join(CHART_DIR, f'qd_bootstrap_ci_{hk}.png'),
                    dpi=150, bbox_inches='tight')
        plt.close(fig)


# ============================================================================
# 10. REPORT
# ============================================================================

def print_report(all_results, friedman_obj, kruskal_results, cochran_result,
                 chi2_conv_results, friedman_wt, wt_kruskal_results,
                 margin_results, scaling_results):

    bench_keys = list(all_results.keys())
    n_bench = len(bench_keys)

    # ── (a) Mean final objective error ──
    print(f"\n{'='*120}")
    print(f"  (a) MEAN FINAL OBJECTIVE ERROR ACROSS {n_bench} BENCHMARK SCENARIOS")
    print("  Friedman test (block design) + Nemenyi post-hoc | Per-scenario: Kruskal-Wallis + Dunn's (Bonferroni)")
    print(f"  alpha = {ALPHA}")
    print(f"{'='*120}")

    print(f"\n  Friedman test (omnibus, {friedman_obj['n_benchmarks']} benchmarks as blocks):")
    print(f"    chi2 = {friedman_obj['chi2']:.4f},  p = {friedman_obj['p_friedman']:.2e},  "
          f"sig = {'YES' if friedman_obj['sig'] else 'NO'}")
    print("    Mean ranks: ", end='')
    for ak in ALGO_KEYS:
        print(f"{ALGO_LABELS[ak]}={friedman_obj['mean_ranks'][ak]:.2f}  ", end='')
    print(f"\n    Best (lowest rank): {ALGO_LABELS[friedman_obj['best_algo']]}")

    if friedman_obj['sig']:
        print(f"\n    Nemenyi post-hoc (CD = {friedman_obj['critical_diff']:.3f}):")
        for np_ in friedman_obj['nemenyi']:
            sig_mark = '*' if np_['significant'] else ' '
            print(f"      {sig_mark} {np_['pair']:>30}: rank_diff={np_['rank_diff']:.3f}  "
                  f"{'SIG' if np_['significant'] else '   '}  {np_['direction']}")

    # Per-scenario Kruskal-Wallis
    print("\n  Per-scenario Kruskal-Wallis (all 4 algorithms):")
    wins = {k: 0 for k in ALGO_KEYS}
    for kr in kruskal_results:
        entry = all_results[kr['key']]
        best_label = ALGO_LABELS.get(kr['best_algo'], kr['best_algo'])
        wins[kr['best_algo']] = wins.get(kr['best_algo'], 0) + 1
        print(f"    {kr['key']:<25} H={kr['h_stat']:>8.2f}  p={kr['p_kruskal']:.2e}  "
              f"{'SIG' if kr['sig'] else '   '}  best={best_label}")

    print("\n  Scenario wins (by mean obj. error): ", end='')
    for ak in ALGO_KEYS:
        print(f"{ALGO_LABELS[ak]}={wins.get(ak, 0)}  ", end='')
    best_variant = max(wins, key=wins.get)
    print(f"\n  >>> Best overall: {ALGO_LABELS[best_variant]}")

    # ── (b) Convergence rate ──
    print(f"\n{'='*120}")
    print(f"  (b) CONVERGENCE RATE ACROSS {N_RUNS} INDEPENDENT RUNS")
    print("  Cochran's Q test (omnibus) + per-scenario chi-square")
    print(f"  alpha = {ALPHA}")
    print(f"{'='*120}")

    print("\n  Cochran's Q test (omnibus):")
    print(f"    Q = {cochran_result['Q_stat']:.4f},  df = {cochran_result['df']},  "
          f"p = {cochran_result['p_value']:.2e},  sig = {'YES' if cochran_result['sig'] else 'NO'}")
    print("    Overall convergence rates:")
    for ak in ALGO_KEYS:
        rate = cochran_result['conv_rates'][ak]
        print(f"      {ALGO_LABELS[ak]:>10}: {rate:.1f}%")

    print(f"\n  Per-scenario chi-square (convergence counts out of {N_RUNS}):")
    for cr in chi2_conv_results:
        counts_str = '  '.join(f"{ALGO_LABELS[ak]}={cr['conv_counts'][ak]}"
                               for ak in ALGO_KEYS)
        print(f"    {cr['key']:<25} chi2={cr['chi2']:>8.2f}  p={cr['p_value']:.2e}  "
              f"{'SIG' if cr['sig'] else '   '}  {counts_str}")

    # Overall convergence summary
    mean_conv = {ak: float(np.mean([all_results[bk][ak]['conv_pct'] for bk in bench_keys]))
                 for ak in ALGO_KEYS}
    best_conv = max(mean_conv, key=mean_conv.get)
    print("\n  Mean convergence rate across all scenarios:")
    for ak in ALGO_KEYS:
        n_100 = sum(1 for bk in bench_keys if all_results[bk][ak]['conv_pct'] == 100.0)
        print(f"    {ALGO_LABELS[ak]:>10}: {mean_conv[ak]:.1f}%  (100% on {n_100}/{n_bench} scenarios)")
    print(f"  >>> Highest mean convergence: {ALGO_LABELS[best_conv]}")

    # ── (c) Wall-clock cost ──
    print(f"\n{'='*120}")
    print("  (c) WALL-CLOCK COST PER RUN")
    print("  Friedman test (block design) + per-scenario Kruskal-Wallis + speedup ratios")
    print(f"  alpha = {ALPHA}")
    print(f"{'='*120}")

    print("\n  Friedman test (omnibus):")
    print(f"    chi2 = {friedman_wt['chi2']:.4f},  p = {friedman_wt['p_friedman']:.2e},  "
          f"sig = {'YES' if friedman_wt['sig'] else 'NO'}")
    print(f"    Fastest overall: {ALGO_LABELS[friedman_wt['fastest']]}")

    print("\n  Grand mean wall-clock time and speedup vs HyGO:")
    for ak in ALGO_KEYS:
        ms = friedman_wt['grand_means_ms'][ak]
        spd = friedman_wt['speedup_vs_hygo'][ak]
        print(f"    {ALGO_LABELS[ak]:>10}: {ms:>8.0f} ms  "
              f"(speedup vs HyGO: {spd:.2f}x)")

    print("\n  Per-scenario Kruskal-Wallis on wall-clock times:")
    for wk in wt_kruskal_results:
        print(f"    {wk['key']:<25} H={wk['h_stat']:>8.2f}  p={wk['p_value']:.2e}  "
              f"{'SIG' if wk['sig'] else '   '}  fastest={ALGO_LABELS[wk['fastest']]}")

    # ── (d) Practically meaningful margin vs HyGO ──
    print(f"\n{'='*120}")
    print("  (d) PRACTICALLY MEANINGFUL MARGIN — EACH HyDE VARIANT vs HyGO")
    print("  Wilcoxon rank-sum + Cliff's delta + bootstrap 95% CI on mean difference")
    print(f"  alpha = {ALPHA}")
    print(f"{'='*120}")

    wins = {k: 0 for k in HYDE_KEYS}
    wins['hygo'] = 0
    ties = 0

    for pr in margin_results:
        sig_mark = '*' if pr['sig'] else ' '
        d_str = f"d={pr['cliffs_delta']:+.3f}({pr['d_magnitude'][:3]})"
        ci_str = f"CI=[{pr['bootstrap_ci_lo']:+.2e}, {pr['bootstrap_ci_hi']:+.2e}]"
        print(f"  {sig_mark} {pr['key']:<22} {pr['hyde_label']:>9} vs HyGO  "
              f"U={pr['u_stat']:>10.0f}  p={pr['p_value']:.2e}  {d_str:>22}  "
              f"{ci_str}  {pr['direction']}")
        if pr['sig']:
            if 'HyGO better' in pr['direction']:
                wins['hygo'] += 1
            else:
                wins[pr['hyde_key']] += 1
        else:
            ties += 1

    print("\n  Significant wins:")
    for k in HYDE_KEYS:
        print(f"    {ALGO_LABELS[k]:>10}: {wins[k]}")
    print(f"    {'HyGO':>10}: {wins['hygo']}")
    print(f"    {'No sig. diff':>10}: {ties}")

    # Cliff's delta summary
    print("\n  Effect size summary (Cliff's delta, thresholds: neg<0.147, sm>=0.147, md>=0.33, lg>=0.474):")
    for hk in HYDE_KEYS:
        hk_results = [pr for pr in margin_results if pr['hyde_key'] == hk]
        ds = [pr['cliffs_delta'] for pr in hk_results]
        mean_d = float(np.mean(ds))
        n_large = sum(1 for d in ds if abs(d) >= 0.474)
        n_med   = sum(1 for d in ds if 0.33 <= abs(d) < 0.474)
        n_sm    = sum(1 for d in ds if 0.147 <= abs(d) < 0.33)
        n_neg   = sum(1 for d in ds if abs(d) < 0.147)
        print(f"    {ALGO_LABELS[hk]:>10} vs HyGO:  mean d={mean_d:>+.3f}  "
              f"large={n_large}  medium={n_med}  small={n_sm}  negligible={n_neg}")

    all_hyde_wins = sum(wins[k] for k in HYDE_KEYS)
    if all_hyde_wins > wins['hygo']:
        verdict = f"HyDE variants superior ({all_hyde_wins} total wins vs HyGO's {wins['hygo']})"
    elif wins['hygo'] > all_hyde_wins:
        verdict = f"HyGO superior ({wins['hygo']} wins vs HyDE's {all_hyde_wins})"
    else:
        verdict = f"No clear winner ({all_hyde_wins} each)"
    print(f"  >>> {verdict}")

    # ── (e) Dimensionality scaling ──
    print(f"\n{'='*120}")
    print("  (e) DIMENSIONALITY SCALING — 2D vs 25D")
    print("  Wilcoxon rank-sum + Cliff's delta + CV + degradation ratio")
    print(f"  {len(SCALABLE_FUNCTIONS)} scalable functions: {', '.join(SCALABLE_FUNCTIONS)}")
    print(f"  alpha = {ALPHA}")
    print(f"{'='*120}")

    for fname in SCALABLE_FUNCTIONS:
        fn_rows = [r for r in scaling_results if r['fname'] == fname]
        if not fn_rows:
            continue
        print(f"\n  {fname.upper()}")
        for row in fn_rows:
            sig_mark = '*' if row['sig'] else ' '
            print(f"    {sig_mark} {row['algo_label']:>10}  "
                  f"U={row['u_stat']:>10.0f}  p={row['p_value']:.2e}  "
                  f"d={row['cliffs_delta']:+.3f}({row['d_magnitude'][:3]})  "
                  f"mean_2D={row['mean_2d']:>10.4e}  mean_25D={row['mean_25d']:>10.4e}  "
                  f"CV_2D={row['cv_2d']:.4f}  CV_25D={row['cv_25d']:.4f}  "
                  f"deg_ratio={row['degradation_ratio']:.2f}  {row['direction']}")

    # Summary: which algorithm degrades least
    print("\n  Consistency at 25D (mean CV) and mean degradation ratio:")
    for ak in ALGO_KEYS:
        ak_rows = [r for r in scaling_results if r['algo_key'] == ak]
        mean_cv25 = float(np.mean([r['cv_25d'] for r in ak_rows])) if ak_rows else float('nan')
        mean_deg  = float(np.mean([r['degradation_ratio'] for r in ak_rows
                                   if np.isfinite(r['degradation_ratio'])])) if ak_rows else float('nan')
        print(f"    {ALGO_LABELS[ak]:>10}: mean CV_25D = {mean_cv25:.4f},  "
              f"mean degradation ratio = {mean_deg:.2f}")
    best_cv_key = min(
        ALGO_KEYS,
        key=lambda ak: float(np.mean([r['cv_25d'] for r in scaling_results
                                       if r['algo_key'] == ak]) or float('inf'))
    )
    print(f"  >>> Most consistent at 25D: {ALGO_LABELS[best_cv_key]}")
    print(f"{'='*120}")


# ============================================================================
# 11. CSV OUTPUT
# ============================================================================

def save_per_run_csv(algo_key, fname, dim, results):
    global_opt = _get_global_opt(fname, dim)
    algo_dir = os.path.join(CSV_DIR, algo_key, f"{fname}_{dim}D")
    os.makedirs(algo_dir, exist_ok=True)

    for i, r in enumerate(results):
        run_file = os.path.join(algo_dir, f"run_{i+1:03d}.csv")
        with open(run_file, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow([
                'algorithm', 'scenario', 'dim', 'run_index',
                'final_best_cost', 'final_objective_error',
                'total_evals', 'convergence_gen', 'wall_clock_ms',
            ])
            obj_err = r['best_cost'] - global_opt
            writer.writerow([
                algo_key, fname, dim, i + 1,
                r['best_cost'], obj_err,
                r['evals'], r.get('conv_gen', ''), r['wall_ms'],
            ])
            writer.writerow([])
            writer.writerow(['evaluation', 'best_cost'])
            for ev_idx, cost in enumerate(r['cost_history'], start=1):
                writer.writerow([ev_idx, cost])


def save_summary_csv(all_results):
    summary_file = os.path.join(CSV_DIR, 'benchmark_summary.csv')
    with open(summary_file, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow([
            'scenario', 'dim', 'algorithm',
            'conv_pct', 'mean_best', 'median_best', 'std_best',
            'min_best', 'max_best', 'iqr_best', 'cv',
            'mean_obj_error', 'std_obj_error',
            'mean_conv_gen', 'std_conv_gen',
            'mean_evals', 'std_evals',
            'mean_wall_ms', 'std_wall_ms',
            'mean_auc', 'std_auc',
        ])
        for bench_key, entry in all_results.items():
            for ak in ALGO_KEYS:
                s = entry[ak]
                writer.writerow([
                    s['fname'], s['dim'], ALGO_LABELS[ak],
                    s['conv_pct'], s['mean_best'], s['median_best'], s['std_best'],
                    s['min_best'], s['max_best'], s['iqr_best'], s['cv'],
                    s['mean_obj_error'], s['std_obj_error'],
                    s['mean_conv_gen'] if s['mean_conv_gen'] is not None else '',
                    s['std_conv_gen']  if s['std_conv_gen']  is not None else '',
                    s['mean_evals'], s['std_evals'],
                    s['mean_wall_ms'], s['std_wall_ms'],
                    s['mean_auc'], s['std_auc'],
                ])


def save_raw_costs_csv(all_results):
    raw_file = os.path.join(CSV_DIR, 'benchmark_raw_costs.csv')
    with open(raw_file, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow([
            'scenario', 'dim', 'algorithm', 'run_index',
            'best_cost', 'wall_ms', 'evals'
        ])
        for bench_key, entry in all_results.items():
            for ak in ALGO_KEYS:
                s = entry[ak]
                for i in range(len(s['raw_costs'])):
                    writer.writerow([
                        s['fname'], s['dim'], ALGO_LABELS[ak], i + 1,
                        s['raw_costs'][i], s['raw_wall_ms'][i], s['raw_evals'][i]
                    ])


def save_qa_csv(friedman_obj, kruskal_results):
    """Save question (a) results."""
    with open(os.path.join(CSV_DIR, 'qa_friedman_objective_error.csv'), 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['test', 'chi2', 'p_value', 'significant', 'best_algo'])
        writer.writerow(['Friedman', friedman_obj['chi2'], friedman_obj['p_friedman'],
                         friedman_obj['sig'], ALGO_LABELS[friedman_obj['best_algo']]])
        writer.writerow([])
        writer.writerow(['pair', 'rank_diff', 'critical_diff', 'significant', 'direction'])
        for np_ in friedman_obj['nemenyi']:
            writer.writerow([np_['pair'], np_['rank_diff'], np_['critical_diff'],
                             np_['significant'], np_['direction']])

    with open(os.path.join(CSV_DIR, 'qa_kruskal_per_scenario.csv'), 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['scenario', 'h_stat', 'p_kruskal', 'significant', 'best_algo'])
        for kr in kruskal_results:
            writer.writerow([kr['key'], kr['h_stat'], kr['p_kruskal'],
                             kr['sig'], ALGO_LABELS.get(kr['best_algo'], kr['best_algo'])])

    with open(os.path.join(CSV_DIR, 'qa_dunns_bonferroni.csv'), 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['scenario', 'pair', 'mean_rank_i', 'mean_rank_j',
                         'z_stat', 'p_bonferroni', 'significant',
                         'cliffs_delta', 'd_magnitude', 'direction'])
        for kr in kruskal_results:
            for ph in kr['posthoc']:
                writer.writerow([kr['key'], ph['pair'],
                                 ph['mean_rank_i'], ph['mean_rank_j'],
                                 ph['z_stat'], ph['p_bonferroni'], ph['significant'],
                                 ph['cliffs_delta'], ph['d_magnitude'], ph['direction']])


def save_qb_csv(cochran_result, chi2_conv_results):
    """Save question (b) results."""
    with open(os.path.join(CSV_DIR, 'qb_convergence_tests.csv'), 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['test', 'Q_stat', 'df', 'p_value', 'significant'])
        writer.writerow(['Cochran_Q', cochran_result['Q_stat'], cochran_result['df'],
                         cochran_result['p_value'], cochran_result['sig']])
        writer.writerow([])
        cols = ['scenario', 'chi2', 'p_value', 'significant'] + \
               [f'conv_{ALGO_LABELS[ak]}' for ak in ALGO_KEYS]
        writer.writerow(cols)
        for cr in chi2_conv_results:
            row = [cr['key'], cr['chi2'], cr['p_value'], cr['sig']]
            row += [cr['conv_counts'][ak] for ak in ALGO_KEYS]
            writer.writerow(row)


def save_qc_csv(friedman_wt, wt_kruskal_results):
    """Save question (c) results."""
    with open(os.path.join(CSV_DIR, 'qc_wall_time_analysis.csv'), 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['test', 'chi2', 'p_value', 'significant', 'fastest'])
        writer.writerow(['Friedman', friedman_wt['chi2'], friedman_wt['p_friedman'],
                         friedman_wt['sig'], ALGO_LABELS[friedman_wt['fastest']]])
        writer.writerow([])
        writer.writerow(['algorithm', 'grand_mean_ms', 'speedup_vs_hygo'])
        for ak in ALGO_KEYS:
            writer.writerow([ALGO_LABELS[ak], friedman_wt['grand_means_ms'][ak],
                             friedman_wt['speedup_vs_hygo'][ak]])
        writer.writerow([])
        writer.writerow(['scenario', 'h_stat', 'p_value', 'significant', 'fastest'])
        for wk in wt_kruskal_results:
            writer.writerow([wk['key'], wk['h_stat'], wk['p_value'],
                             wk['sig'], ALGO_LABELS[wk['fastest']]])


def save_qd_csv(margin_results):
    """Save question (d) results."""
    with open(os.path.join(CSV_DIR, 'qd_margin_vs_hygo.csv'), 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['scenario', 'hyde_variant', 'u_stat', 'p_value', 'significant',
                         'hyde_mean', 'hygo_mean', 'mean_diff',
                         'cliffs_delta', 'd_magnitude',
                         'bootstrap_ci_lo', 'bootstrap_ci_hi', 'direction'])
        for pr in margin_results:
            writer.writerow([pr['key'], pr['hyde_label'], pr['u_stat'], pr['p_value'],
                             pr['sig'], pr['hyde_mean'], pr['hygo_mean'], pr['mean_diff'],
                             pr['cliffs_delta'], pr['d_magnitude'],
                             pr['bootstrap_ci_lo'], pr['bootstrap_ci_hi'], pr['direction']])


def save_qe_csv(scaling_results):
    """Save question (e) results."""
    with open(os.path.join(CSV_DIR, 'qe_scaling_analysis.csv'), 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['scenario', 'algorithm', 'u_stat', 'p_value', 'significant',
                         'mean_2d', 'mean_25d', 'cv_2d', 'cv_25d',
                         'cliffs_delta', 'd_magnitude',
                         'degradation_ratio', 'direction'])
        for row in scaling_results:
            writer.writerow([row['fname'], row['algo_label'], row['u_stat'],
                             row['p_value'], row['sig'],
                             row['mean_2d'], row['mean_25d'],
                             row['cv_2d'], row['cv_25d'],
                             row['cliffs_delta'], row['d_magnitude'],
                             row['degradation_ratio'], row['direction']])


# ============================================================================
# 12. DOCX REPORT
# ============================================================================

def generate_docx_report(all_results, friedman_obj, kruskal_results,
                         cochran_result, chi2_conv_results,
                         friedman_wt, margin_results, scaling_results):
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as e:
        print(f"  [skip DOCX report — missing dependency: {e}]")
        return

    BLACK = RGBColor(0, 0, 0)

    docx_path = os.path.join(HERE, 'benchmark_report.docx')
    doc = Document()

    # ── Global font defaults: Times New Roman 12pt, black ──
    style = doc.styles['Normal']
    style.font.name  = 'Times New Roman'
    style.font.size  = Pt(12)
    style.font.color.rgb = BLACK
    pf = style.paragraph_format
    pf.line_spacing  = 2.0
    pf.space_before  = Pt(0)
    pf.space_after   = Pt(0)

    for hname in ('Heading 1', 'Heading 2', 'Heading 3'):
        hs = doc.styles[hname]
        hs.font.name      = 'Times New Roman'
        hs.font.size      = Pt(12)
        hs.font.bold      = True
        hs.font.color.rgb = BLACK
        hs.paragraph_format.line_spacing  = 2.0
        hs.paragraph_format.space_before  = Pt(0)
        hs.paragraph_format.space_after   = Pt(0)

    # ── Counters ──
    # Tables: section-based numbering starting at 3.1
    # Figures: sequential starting at 4
    tbl_counter = [0]   # mutable so closures can increment
    fig_counter = [3]   # next figure will be 4

    # ── Helpers ──

    def _para(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing  = 2.0
        p.paragraph_format.space_before  = Pt(0)
        p.paragraph_format.space_after   = Pt(0)
        run = p.add_run(text)
        run.font.name      = 'Times New Roman'
        run.font.size      = Pt(12)
        run.font.color.rgb = BLACK
        run.bold = bold
        return p

    def _heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.line_spacing  = 2.0
        h.paragraph_format.space_before  = Pt(0)
        h.paragraph_format.space_after   = Pt(0)
        for run in h.runs:
            run.font.name      = 'Times New Roman'
            run.font.size      = Pt(12)
            run.font.color.rgb = BLACK
        return h

    def _set_cell_border(cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        for edge, attrs in kwargs.items():
            el = parse_xml(
                f'<w:{edge} {nsdecls("w")} '
                f'w:val="{attrs.get("val","single")}" '
                f'w:sz="{attrs.get("sz","4")}" '
                f'w:space="0" '
                f'w:color="{attrs.get("color","000000")}"/>')
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _table_caption(label_text, description_text):
        """Add a caption paragraph ABOVE the table/figure.
        Label (e.g. 'Table 3.1') is bold, description is italic non-bold."""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing  = 2.0
        p.paragraph_format.space_before  = Pt(0)
        p.paragraph_format.space_after   = Pt(0)
        rl = p.add_run(label_text)
        rl.font.name      = 'Times New Roman'
        rl.font.size      = Pt(12)
        rl.font.color.rgb = BLACK
        rl.bold   = True
        rl.italic = False
        rd = p.add_run(' ' + description_text)
        rd.font.name      = 'Times New Roman'
        rd.font.size      = Pt(12)
        rd.font.color.rgb = BLACK
        rd.bold   = False
        rd.italic = True
        return p

    def _next_table_label():
        tbl_counter[0] += 1
        return f"Table 3.{tbl_counter[0]}"

    def _next_fig_label():
        fig_counter[0] += 1
        return f"Figure {fig_counter[0]}"

    def _add_table(headers, rows, description):
        """Caption on top, then borderless table with header/bottom rules."""
        _table_caption(_next_table_label(), description)

        n_cols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Remove all default borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'</w:tblBorders>')
        for existing in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(existing)
        tblPr.append(borders)

        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = ''
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.line_spacing  = 1.0
            p.paragraph_format.space_before  = Pt(0)
            p.paragraph_format.space_after   = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(h))
            run.font.name      = 'Times New Roman'
            run.font.size      = Pt(12)
            run.font.color.rgb = BLACK
            run.bold = True
            _set_cell_border(hdr_cells[i],
                             top={"val": "single", "sz": "4"},
                             bottom={"val": "single", "sz": "4"})

        # Data rows
        for ri, row_data in enumerate(rows):
            row_cells = table.rows[ri + 1].cells
            is_last = (ri == len(rows) - 1)
            for ci, val in enumerate(row_data):
                row_cells[ci].text = ''
                p = row_cells[ci].paragraphs[0]
                p.paragraph_format.line_spacing  = 1.0
                p.paragraph_format.space_before  = Pt(0)
                p.paragraph_format.space_after   = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(val))
                run.font.name      = 'Times New Roman'
                run.font.size      = Pt(12)
                run.font.color.rgb = BLACK
                if is_last:
                    _set_cell_border(row_cells[ci],
                                     bottom={"val": "single", "sz": "4"})

        # Minimal cell margins
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = parse_xml(
                    f'<w:tcMar {nsdecls("w")}>'
                    f'  <w:top w:w="0" w:type="dxa"/>'
                    f'  <w:left w:w="28" w:type="dxa"/>'
                    f'  <w:bottom w:w="0" w:type="dxa"/>'
                    f'  <w:right w:w="28" w:type="dxa"/>'
                    f'</w:tcMar>')
                for existing in tcPr.findall(qn('w:tcMar')):
                    tcPr.remove(existing)
                tcPr.append(tcMar)

        return table

    def _add_chart(chart_path, description, width_inches=6.0):
        """Caption on top (Figure N + italic description), then image."""
        if not os.path.exists(chart_path):
            return
        _table_caption(_next_fig_label(), description)
        doc.add_picture(chart_path, width=Inches(width_inches))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_p.paragraph_format.space_before = Pt(0)
        last_p.paragraph_format.space_after  = Pt(0)

    bench_keys = list(all_results.keys())

    # ═══════════════════════════════════════════════════════════════════
    # TITLE
    # ═══════════════════════════════════════════════════════════════════
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.line_spacing  = 2.0
    tp.paragraph_format.space_before  = Pt(0)
    tp.paragraph_format.space_after   = Pt(0)
    run = tp.add_run("HyDE Benchmark Report")
    run.font.name      = 'Times New Roman'
    run.font.size      = Pt(12)
    run.font.color.rgb = BLACK
    run.bold = True

    _para(
        f"Conditions: {len(TEST_CASES)} benchmark scenarios, "
        f"{N_RUNS} independent runs, {MAX_EVALS} evaluation budget, "
        f"alpha = {ALPHA}.")
    _para(
        "Algorithms: HyDE-bin, HyDE-qub, HyDE-con, HyGO.")
    _para(
        "(a) Friedman + Nemenyi + per-scenario Kruskal-Wallis + "
        "Dunn\u2019s (Bonferroni) + Cliff\u2019s delta. "
        "(b) Cochran\u2019s Q + per-scenario chi-square. "
        "(c) Friedman on wall times + Kruskal-Wallis + speedup ratios. "
        "(d) Wilcoxon rank-sum + Cliff\u2019s delta + bootstrap 95% CI. "
        "(e) Wilcoxon + Cliff\u2019s delta + CV + degradation ratio.")

    from datetime import datetime
    _para(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # (a) MEAN FINAL OBJECTIVE ERROR
    # ═══════════════════════════════════════════════════════════════════
    _heading("(a) Mean Final Objective Error")

    n_sig_kruskal = sum(1 for kr in kruskal_results if kr['sig'])
    wins_a = {k: 0 for k in ALGO_KEYS}
    for kr in kruskal_results:
        if kr['best_algo'] in wins_a:
            wins_a[kr['best_algo']] += 1
    best_a = max(wins_a, key=wins_a.get)

    _para(
        f"The Friedman test (block design, {friedman_obj['n_benchmarks']} benchmarks) "
        f"yielded \u03c7\u00b2 = {friedman_obj['chi2']:.2f}, "
        f"p = {friedman_obj['p_friedman']:.2e} "
        f"({'significant' if friedman_obj['sig'] else 'not significant'}). "
        f"{ALGO_LABELS[friedman_obj['best_algo']]} achieved the lowest mean rank "
        f"({friedman_obj['mean_ranks'][friedman_obj['best_algo']]:.2f}). "
        f"Per-scenario Kruskal-Wallis tests found significant differences on "
        f"{n_sig_kruskal} of {len(kruskal_results)} scenarios. "
        f"By mean objective error, {ALGO_LABELS[best_a]} won "
        f"{wins_a[best_a]} scenarios.")

    _add_table(
        ['Pair', 'Rank Diff', 'CD', 'Significant', 'Direction'],
        [[np_['pair'], f"{np_['rank_diff']:.3f}",
          f"{np_['critical_diff']:.3f}",
          'Yes' if np_['significant'] else 'No',
          np_['direction']] for np_ in friedman_obj['nemenyi']],
        description="Nemenyi post-hoc pairwise comparisons following Friedman test.")

    _add_table(
        ['Scenario', 'H-stat', 'p-value', 'Sig.', 'Best'],
        [[kr['key'], f"{kr['h_stat']:.2f}", f"{kr['p_kruskal']:.2e}",
          'Yes' if kr['sig'] else 'No',
          ALGO_LABELS.get(kr['best_algo'], kr['best_algo'])]
         for kr in kruskal_results],
        description="Per-scenario Kruskal-Wallis results across all four algorithms.")

    _add_chart(os.path.join(CHART_DIR, 'qa_objective_error_wins.png'),
               "Best algorithm per scenario by mean objective error.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # (b) CONVERGENCE RATE
    # ═══════════════════════════════════════════════════════════════════
    _heading("(b) Convergence Rate")

    mean_conv = {ak: float(np.mean([all_results[bk][ak]['conv_pct']
                                     for bk in bench_keys]))
                 for ak in ALGO_KEYS}
    best_conv = max(mean_conv, key=mean_conv.get)

    _para(
        f"Cochran\u2019s Q test yielded Q = {cochran_result['Q_stat']:.2f}, "
        f"p = {cochran_result['p_value']:.2e} "
        f"({'significant' if cochran_result['sig'] else 'not significant'}), "
        f"indicating that convergence rates "
        f"{'differ significantly' if cochran_result['sig'] else 'do not differ significantly'} "
        f"across algorithms. {ALGO_LABELS[best_conv]} achieved the highest mean "
        f"convergence rate ({mean_conv[best_conv]:.1f}%).")

    conv_rows = []
    for cr in chi2_conv_results:
        row = [cr['key']]
        row += [f"{cr['conv_counts'][ak]}/{N_RUNS}" for ak in ALGO_KEYS]
        row += [f"{cr['chi2']:.2f}", f"{cr['p_value']:.2e}",
                'Yes' if cr['sig'] else 'No']
        conv_rows.append(row)
    _add_table(
        ['Scenario'] + [ALGO_LABELS[ak] for ak in ALGO_KEYS] + ['\u03c7\u00b2', 'p', 'Sig.'],
        conv_rows,
        description="Per-scenario convergence counts and chi-square test.")

    _add_chart(os.path.join(CHART_DIR, 'qb_convergence_rate.png'),
               f"Convergence rate per benchmark across {N_RUNS} independent runs.")

    for ak in ALGO_KEYS:
        n_100 = sum(1 for bk in bench_keys
                    if all_results[bk][ak]['conv_pct'] == 100.0)
        _para(f"{ALGO_LABELS[ak]}: mean convergence = {mean_conv[ak]:.1f}%, "
              f"100% convergence on {n_100}/{len(bench_keys)} scenarios.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # (c) WALL-CLOCK COST PER RUN
    # ═══════════════════════════════════════════════════════════════════
    _heading("(c) Wall-Clock Cost per Run")

    wt_text = (
        f"The Friedman test on median wall-clock times yielded "
        f"\u03c7\u00b2 = {friedman_wt['chi2']:.2f}, "
        f"p = {friedman_wt['p_friedman']:.2e} "
        f"({'significant' if friedman_wt['sig'] else 'not significant'}). "
        f"{ALGO_LABELS[friedman_wt['fastest']]} was the fastest overall "
        f"({friedman_wt['grand_means_ms'][friedman_wt['fastest']]:.0f} ms mean). ")
    for ak in ALGO_KEYS:
        if ak != 'hygo':
            spd = friedman_wt['speedup_vs_hygo'][ak]
            wt_text += (
                f"{ALGO_LABELS[ak]} is {spd:.2f}x "
                f"{'faster' if spd > 1 else 'slower'} than HyGO. ")
    _para(wt_text)

    _add_table(
        ['Algorithm', 'Grand Mean (ms)', 'Speedup vs HyGO'],
        [[ALGO_LABELS[ak],
          f"{friedman_wt['grand_means_ms'][ak]:.0f}",
          f"{friedman_wt['speedup_vs_hygo'][ak]:.2f}x"]
         for ak in ALGO_KEYS],
        description="Wall-clock cost summary with speedup ratios relative to HyGO.")

    _add_chart(os.path.join(CHART_DIR, 'qc_wall_time_summary.png'),
               "Average wall-clock time across all benchmarks.")
    _add_chart(os.path.join(CHART_DIR, 'qc_wall_time_per_benchmark.png'),
               "Mean wall-clock time per benchmark scenario.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # (d) PRACTICALLY MEANINGFUL MARGIN vs HyGO
    # ═══════════════════════════════════════════════════════════════════
    _heading("(d) Practically Meaningful Margin vs HyGO")

    wins_d = {k: 0 for k in HYDE_KEYS}
    wins_d['hygo'] = 0
    ties_d = 0
    for pr in margin_results:
        if pr['sig']:
            if 'HyGO better' in pr['direction']:
                wins_d['hygo'] += 1
            else:
                wins_d[pr['hyde_key']] += 1
        else:
            ties_d += 1
    all_hyde_wins = sum(wins_d[k] for k in HYDE_KEYS)

    d_text = (
        f"Wilcoxon rank-sum tests with Cliff\u2019s delta and bootstrap 95% CIs were "
        f"applied for each HyDE variant against HyGO across all {len(TEST_CASES)} scenarios. "
        f"The HyDE variants collectively achieved {all_hyde_wins} significant wins "
        f"vs HyGO\u2019s {wins_d['hygo']}; {ties_d} showed no significant difference. ")
    for hk in HYDE_KEYS:
        hk_r = [pr for pr in margin_results if pr['hyde_key'] == hk]
        ds = [pr['cliffs_delta'] for pr in hk_r]
        mean_d = float(np.mean(ds))
        n_lg = sum(1 for d in ds if abs(d) >= 0.474)
        d_text += (
            f"{ALGO_LABELS[hk]}: mean Cliff\u2019s d = {mean_d:+.3f} "
            f"({interpret_cliffs_delta(mean_d)}), {n_lg} large effects. ")
    _para(d_text)

    _add_table(
        ['Scenario', 'Variant', 'U', 'p', 'Sig.', "Cliff\u2019s d",
         'CI lo', 'CI hi', 'Direction'],
        [[pr['key'], pr['hyde_label'],
          f"{pr['u_stat']:.0f}", f"{pr['p_value']:.2e}",
          'Yes' if pr['sig'] else 'No',
          f"{pr['cliffs_delta']:+.3f}",
          f"{pr['bootstrap_ci_lo']:+.2e}",
          f"{pr['bootstrap_ci_hi']:+.2e}",
          pr['direction']] for pr in margin_results],
        description="Wilcoxon rank-sum + Cliff\u2019s delta + bootstrap 95% CI "
                    "for each HyDE variant vs HyGO.")

    es_rows = []
    for hk in HYDE_KEYS:
        hk_r = [pr for pr in margin_results if pr['hyde_key'] == hk]
        ds = [pr['cliffs_delta'] for pr in hk_r]
        mean_d = float(np.mean(ds))
        es_rows.append([
            f"{ALGO_LABELS[hk]} vs HyGO",
            f"{mean_d:+.3f}",
            str(sum(1 for d in ds if abs(d) >= 0.474)),
            str(sum(1 for d in ds if 0.33 <= abs(d) < 0.474)),
            str(sum(1 for d in ds if 0.147 <= abs(d) < 0.33)),
            str(sum(1 for d in ds if abs(d) < 0.147)),
        ])
    _add_table(
        ['Comparison', 'Mean d', 'Large', 'Medium', 'Small', 'Negligible'],
        es_rows,
        description="Cliff\u2019s delta effect size distribution.")

    _add_chart(os.path.join(CHART_DIR, 'qd_margin_wins.png'),
               "Significant wins from Wilcoxon rank-sum test, each HyDE variant vs HyGO.")
    for hk in HYDE_KEYS:
        _add_chart(os.path.join(CHART_DIR, f'qd_bootstrap_ci_{hk}.png'),
                   f"Bootstrap 95% CI on mean difference: {ALGO_LABELS[hk]} vs HyGO.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # (e) DIMENSIONALITY SCALING: 2D vs 25D
    # ═══════════════════════════════════════════════════════════════════
    _heading("(e) Dimensionality Scaling: 2D vs 25D")

    if scaling_results:
        e_parts = []
        for ak in ALGO_KEYS:
            ak_rows = [r for r in scaling_results if r['algo_key'] == ak]
            mean_cv25 = float(np.mean([r['cv_25d'] for r in ak_rows]))
            n_deg = sum(1 for r in ak_rows
                        if r['sig'] and r['direction'] == 'degraded')
            finite_degs = [r['degradation_ratio'] for r in ak_rows
                           if np.isfinite(r['degradation_ratio'])]
            mean_deg = float(np.mean(finite_degs)) if finite_degs else float('nan')
            e_parts.append(
                f"{ALGO_LABELS[ak]}: significant degradation on "
                f"{n_deg}/{len(ak_rows)} functions, "
                f"mean CV at 25D = {mean_cv25:.4f}, "
                f"mean degradation ratio = {mean_deg:.2f}.")

        best_cv_key = min(
            ALGO_KEYS,
            key=lambda ak: float(np.mean([r['cv_25d'] for r in scaling_results
                                           if r['algo_key'] == ak]) or float('inf')))
        e_parts.append(
            f"Most consistent at 25D: {ALGO_LABELS[best_cv_key]}.")

        _para(' '.join(e_parts))

        _add_table(
            ['Function', 'Algorithm', 'U', 'p', 'Sig.', "d",
             'CV 2D', 'CV 25D', 'Deg. Ratio', 'Dir.'],
            [[row['fname'], row['algo_label'],
              f"{row['u_stat']:.0f}", f"{row['p_value']:.2e}",
              'Yes' if row['sig'] else 'No',
              f"{row['cliffs_delta']:+.3f}",
              f"{row['cv_2d']:.4f}", f"{row['cv_25d']:.4f}",
              f"{row['degradation_ratio']:.2f}",
              row['direction']] for row in scaling_results],
            description="Dimensionality scaling analysis with degradation ratio.")

        _add_chart(os.path.join(CHART_DIR, 'qe_cv_25d.png'),
                   "Coefficient of Variation at 25D per algorithm for each scalable benchmark.")
        _add_chart(os.path.join(CHART_DIR, 'qe_degradation_heatmap.png'),
                   "Degradation ratio heatmap: mean_25D / mean_2D (lower = scales better).")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # PER-BENCHMARK RESULTS
    # ═══════════════════════════════════════════════════════════════════
    _heading("Per-Benchmark Results")
    _para(
        f"Mean convergence curve and final cost distribution for each of the "
        f"{len(bench_keys)} benchmark scenarios.")

    for bi, bench_key in enumerate(bench_keys):
        entry = all_results[bench_key]
        fname = entry['hyde_bin']['fname']
        dim   = entry['hyde_bin']['dim']

        _heading(f"{fname.upper()} {dim}D", level=2)

        means = {ak: entry[ak]['mean_best'] for ak in ALGO_KEYS}
        convs = {ak: entry[ak]['conv_pct']  for ak in ALGO_KEYS}
        best_h = min(means, key=means.get)
        best_c = max(convs, key=convs.get)

        _para(
            f"Best mean cost: {ALGO_LABELS[best_h]} ({means[best_h]:.4e}). "
            f"Best convergence: {ALGO_LABELS[best_c]} ({convs[best_c]:.1f}%). "
            f"Costs: "
            + ", ".join(f"{ALGO_LABELS[ak]} = {means[ak]:.4e}"
                        for ak in ALGO_KEYS)
            + ".")

        _add_chart(os.path.join(CHART_DIR, f'{bench_key}.png'),
                   f"{fname.upper()} {dim}D convergence curve and cost distribution.")

        _add_table(
            ['Algorithm', 'Conv %', 'Mean Best', 'Std', 'CV',
             'Mean Gen', 'Wall (ms)'],
            [[ALGO_LABELS[ak],
              f"{entry[ak]['conv_pct']:.1f}",
              f"{entry[ak]['mean_best']:.4e}",
              f"{entry[ak]['std_best']:.2e}",
              f"{entry[ak]['cv']:.4f}",
              str(entry[ak]['mean_conv_gen']) if entry[ak]['mean_conv_gen'] else '-',
              f"{entry[ak]['mean_wall_ms']:.0f}"]
             for ak in ALGO_KEYS],
            description=f"{fname.upper()} {dim}D summary statistics.")

        if bi < len(bench_keys) - 1:
            doc.add_page_break()

    doc.save(docx_path)
    print(f"  DOCX report saved to: {docx_path}")


# ============================================================================
# 13. MAIN
# ============================================================================

def main():
    from hyde_bin import HyDEBin
    from hyde_con import HyDECon
    from hyde_qub import HyDEQub
    from hygo import HyGO

    print(f"\n{chr(0x2554)+chr(0x2550)*118+chr(0x2557)}")
    print(f"{chr(0x2551)} {'BENCHMARK: HyDE-bin  |  HyDE-qub  |  HyDE-con  |  HyGO':^116} {chr(0x2551)}")
    print(f"{chr(0x2551)} {f'{N_RUNS} runs  x  {len(TEST_CASES)} benchmarks  x  4 algorithms = {N_RUNS*len(TEST_CASES)*4:,} total runs  |  Max evals: {MAX_EVALS}  |  alpha = {ALPHA}':^116} {chr(0x2551)}")
    print(f"{chr(0x2551)} {'(a) Friedman+Nemenyi  (b) Cochran Q  (c) Friedman wall-time  (d) Wilcoxon+bootstrap CI  (e) Scaling':^116} {chr(0x2551)}")
    print(f"{chr(0x255a)+chr(0x2550)*118+chr(0x255d)}\n")

    hyde_kwargs = dict(pop_size=None, max_gen=50, phase_split=0.60)
    hygo_kwargs = dict(Nb=12, NG=50, Nexplor=70, Nexploit=30,
                       Ne=1, ps=0.5, Pc=0.55, Pm=0.45, Pr=0.0)

    all_results     = {}
    margin_stats    = []
    t_total = time.time()

    for fname, dim in TEST_CASES:
        key = f"{fname}_{dim}D"
        tag = f"{fname.upper()} {dim}D"

        t0 = time.time()

        bin_res  = run_case(HyDEBin, fname, dim, {**hyde_kwargs, 'Nb': 12}, N_RUNS, seed_base=0)
        qub_res  = run_case(HyDEQub, fname, dim, hyde_kwargs,               N_RUNS, seed_base=0)
        con_res  = run_case(HyDECon, fname, dim, hyde_kwargs,               N_RUNS, seed_base=0)
        hkw      = {**hygo_kwargs, 'NT': 7 if dim <= 5 else 100}
        hygo_res = run_case(HyGO,    fname, dim, hkw,                       N_RUNS, seed_base=0)

        elapsed = time.time() - t0

        save_per_run_csv('hyde_bin', fname, dim, bin_res)
        save_per_run_csv('hyde_qub', fname, dim, qub_res)
        save_per_run_csv('hyde_con', fname, dim, con_res)
        save_per_run_csv('hygo',     fname, dim, hygo_res)

        entry = {
            'hyde_bin': summarize(bin_res,  fname, dim),
            'hyde_qub': summarize(qub_res,  fname, dim),
            'hyde_con': summarize(con_res,  fname, dim),
            'hygo':     summarize(hygo_res, fname, dim),
        }
        all_results[key] = entry

        medians = {k: entry[k]['median_best'] for k in ALGO_KEYS}
        best_k  = min(medians, key=medians.get)

        print(
            f"  {tag:<22}"
            f"  BIN={entry['hyde_bin']['median_best']:>10.3e}"
            f"  QUB={entry['hyde_qub']['median_best']:>10.3e}"
            f"  CON={entry['hyde_con']['median_best']:>10.3e}"
            f"  |  HyGO={entry['hygo']['median_best']:>10.3e}"
            f"  |  best={ALGO_LABELS[best_k]}"
            f"  [{elapsed:.1f}s]"
        )

        # (d) Margin vs HyGO
        for hyde_key in HYDE_KEYS:
            margin_stats.append(wilcoxon_margin_vs_hygo(key, entry, hyde_key))

        # Incremental JSON cache
        with open(os.path.join(HERE, 'benchmark_results.json'), 'w') as f:
            json.dump(all_results, f, indent=2)

    # ── Statistical analyses ──

    # (a) Friedman + per-scenario Kruskal-Wallis
    friedman_obj = friedman_objective_error(all_results)
    kruskal_results = [run_kruskal_per_scenario(key, entry)
                       for key, entry in all_results.items()]

    # (b) Cochran's Q + per-scenario chi-square
    cochran_result = cochrans_q_test(all_results)
    chi2_conv_results = [chi2_convergence_per_scenario(key, entry)
                         for key, entry in all_results.items()]

    # (c) Friedman wall time + per-scenario Kruskal-Wallis
    friedman_wt = friedman_wall_time(all_results)
    wt_kruskal_results = [kruskal_wall_time_per_scenario(key, entry)
                          for key, entry in all_results.items()]

    # (e) Scaling analysis
    scaling_results = run_scaling_analysis(all_results)

    # ── Reports ──
    print_report(all_results, friedman_obj, kruskal_results, cochran_result,
                 chi2_conv_results, friedman_wt, wt_kruskal_results,
                 margin_stats, scaling_results)

    # ── Charts ──
    make_charts(all_results, kruskal_results, margin_stats)
    make_scaling_chart(scaling_results)
    make_cost_charts(all_results)
    make_convergence_charts(all_results)
    make_figure5_curves(all_results)
    make_bootstrap_ci_chart(margin_stats)

    # ── DOCX ──
    generate_docx_report(all_results, friedman_obj, kruskal_results,
                         cochran_result, chi2_conv_results,
                         friedman_wt, margin_stats, scaling_results)

    # ── CSV outputs ──
    save_summary_csv(all_results)
    save_raw_costs_csv(all_results)
    save_qa_csv(friedman_obj, kruskal_results)
    save_qb_csv(cochran_result, chi2_conv_results)
    save_qc_csv(friedman_wt, wt_kruskal_results)
    save_qd_csv(margin_stats)
    save_qe_csv(scaling_results)

    # Final JSON
    with open(os.path.join(HERE, 'benchmark_results.json'), 'w') as f:
        json.dump(all_results, f, indent=2)

    print(f"\n  Total time: {time.time()-t_total:.1f}s")
    print(f"  CSV data saved to:  {CSV_DIR}/")
    print(f"    Per-run data:     {CSV_DIR}/{{algorithm}}/{{scenario}}_{{dim}}D/run_XXX.csv")
    print(f"    Summary:          {CSV_DIR}/benchmark_summary.csv")
    print(f"    Raw costs:        {CSV_DIR}/benchmark_raw_costs.csv")
    print(f"    (a) Friedman:     {CSV_DIR}/qa_friedman_objective_error.csv")
    print(f"    (a) Kruskal:      {CSV_DIR}/qa_kruskal_per_scenario.csv")
    print(f"    (a) Dunn's:       {CSV_DIR}/qa_dunns_bonferroni.csv")
    print(f"    (b) Convergence:  {CSV_DIR}/qb_convergence_tests.csv")
    print(f"    (c) Wall time:    {CSV_DIR}/qc_wall_time_analysis.csv")
    print(f"    (d) Margin:       {CSV_DIR}/qd_margin_vs_hygo.csv")
    print(f"    (e) Scaling:      {CSV_DIR}/qe_scaling_analysis.csv")
    print(f"  Charts saved to:    {CHART_DIR}/")
    print(f"  DOCX report:        {HERE}/benchmark_report.docx")


if __name__ == '__main__':
    main()
